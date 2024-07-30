import streamlit as st
from docx import Document
from datetime import datetime

# Define the questions for the MoU
mou_questions = [
    "Enter the date of the MoU:",
    "Enter the first party's name:",
    "Enter the first party's address (Line 1):",
    "Enter the first party's address (Line 2):",
    "Enter the first party's city, state, and pin code:",
    "Enter the second party's name:",
    "Enter the second party's address (Line 1):",
    "Enter the second party's address (Line 2):",
    "Enter the second party's city, state, and pin code:",
    "Enter the purpose of the MoU:",
    "Enter the duration of the MoU:",
    "Enter any specific terms and conditions:"
]

def create_mou(answers):
    doc = Document()
    
    doc.add_heading('Memorandum of Understanding (MoU)', 0)

    doc.add_paragraph(f"This Memorandum of Understanding (MoU) is made on {answers[0]} between {answers[1]}, residing at {answers[2]}, {answers[3]}, {answers[4]}, hereinafter referred to as the 'First Party' and {answers[5]}, residing at {answers[6]}, {answers[7]}, {answers[8]}, hereinafter referred to as the 'Second Party'.")

    doc.add_paragraph(f"WHEREAS the First Party and the Second Party wish to collaborate on {answers[9]} for a duration of {answers[10]}.")

    doc.add_paragraph(f"NOW THEREFORE, the Parties agree to the following terms and conditions:")

    doc.add_paragraph(f"1. Purpose: The purpose of this MoU is to {answers[9]}.")
    doc.add_paragraph(f"2. Duration: This MoU shall remain in effect for {answers[10]}.")
    doc.add_paragraph(f"3. Terms and Conditions: {answers[11]}")

    return doc

# Streamlit UI for MoU
st.title("MoU Document Generator")
st.write("Answer the following questions to generate your MoU:")

# Initialize session state to keep track of questions and answers
if 'current_mou_question_index' not in st.session_state:
    st.session_state.current_mou_question_index = 0
if 'mou_answers' not in st.session_state:
    st.session_state.mou_answers = {}

# Display the current question
if st.session_state.current_mou_question_index < len(mou_questions):
    question = mou_questions[st.session_state.current_mou_question_index]
    
    if "date" in question.lower():
        answer = st.date_input(question)
    else:
        answer = st.text_input(question)
    
    if st.button("Submit"):
        if answer:
            st.session_state.mou_answers[st.session_state.current_mou_question_index] = answer
            st.session_state.current_mou_question_index += 1
            st.experimental_rerun()

    if st.button("Back") and st.session_state.current_mou_question_index > 0:
        st.session_state.current_mou_question_index -= 1
        st.experimental_rerun()
    
else:
    st.write("Thank you! Generating your MoU...")

    answers = [st.session_state.mou_answers.get(i, "") for i in range(len(mou_questions))]
    document = create_mou(answers)

    # Save the document
    document_path = "MoU.docx"
    document.save(document_path)

    st.write("Your MoU has been generated.")
    with open(document_path, "rb") as file:
        st.download_button("Download MoU", file, file_name=document_path)

    if st.button("Summary"):
        st.write("Summary of your inputs:")
        for i, answer in st.session_state.mou_answers.items():
            st.write(f"{mou_questions[i]} {answer}")
