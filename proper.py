import streamlit as st
from docx import Document

# Function to create rental agreement document
def create_rental_agreement(answers):
    doc = Document()
    
    doc.add_heading('Residential Rental Agreement', 0)

    doc.add_paragraph(f"This agreement made at {answers[0]} on this {answers[1]} between {answers[2]}, residing at {answers[3]}, {answers[4]}, {answers[5]}, hereinafter referred to as the 'LESSOR' of the One Part AND {answers[6]}, residing at {answers[7]}, {answers[8]}, {answers[9]}, hereinafter referred to as the 'LESSEE' of the other Part;")

    doc.add_paragraph(f"WHEREAS the Lessor is the lawful owner of, and otherwise well sufficiently entitled to {answers[10]}, {answers[11]}, {answers[12]} falling in the category, {answers[13]} and comprising of {answers[14]} Bedrooms, {answers[15]} Bathrooms, {answers[16]} Carparks with an extent of {answers[17]} Square Feet hereinafter referred to as the 'said premises';")

    doc.add_paragraph(f"AND WHEREAS at the request of the Lessee, the Lessor has agreed to let the said premises to the tenant for a term of {answers[18]} commencing from {answers[19]} in the manner hereinafter appearing.")

    doc.add_paragraph(f"NOW THIS AGREEMENT WITNESSETH AND IT IS HEREBY AGREED BY AND BETWEEN THE PARTIES AS UNDER:")

    doc.add_paragraph(f"1. That the Lessor hereby grants to the Lessee, the right to enter into and use and remain in the said premises along with the existing fixtures and fittings listed in Annexure 1 to this Agreement and that the Lessee shall be entitled to peacefully possess, and enjoy possession of the said premises, and the other rights herein.")
    doc.add_paragraph(f"2. That the lease hereby granted shall, unless cancelled earlier under any provision of this Agreement, remain in force for a period of {answers[18]}.")
    doc.add_paragraph(f"3. That the Lessee will have the option to terminate this lease by giving {answers[22]} in writing to the Lessor.")
    doc.add_paragraph(f"4. That the Lessee shall have no right to create any sub-lease or assign or transfer in any manner the lease or give to any one the possession of the said premises or any part thereof.")
    doc.add_paragraph(f"5. That the Lessee shall use the said premises only for residential purposes.")
    doc.add_paragraph(f"6. That the Lessor shall, before handing over the said premises, ensure the working of sanitary, electrical and water supply connections and other fittings pertaining to the said premises. It is agreed that it shall be the responsibility of the Lessor for their return in the working condition at the time of re-possession of the said premises (reasonable wear and tear and loss or damage by fire, flood, rains, accident, irresistible force or act of God excepted).")
    doc.add_paragraph(f"7. That the Lessee is not authorized to make any alteration in the construction of the said premises. The Lessee may however install and remove his own fittings and fixtures, provided this is done without causing any excessive damage or loss to the said premises.")
    doc.add_paragraph(f"8. That the day to day repair jobs such as fuse blow out, replacement of light bulbs/tubes, leakage of water taps, maintenance of the water pump and other minor repairs, etc., shall be effected by the Lessee at its own cost, and any major repairs, either structural or to the electrical or water connection, plumbing leaks, water seepage shall be attended to by the Lessor. In the event of the Lessor failing to carry out the repairs on receiving notice from the Lessee, the Lessee shall undertake the necessary repairs and the Lessor will be liable to immediately reimburse costs incurred by the Lessee.")
    doc.add_paragraph(f"9. That the Lessor or its duly authorized agent shall have the right to enter into or upon the said premises or any part thereof at a mutually arranged convenient time for the purpose of inspection.")
    doc.add_paragraph(f"10. That the Lessee shall use the said premises along with its fixtures and fitting in careful and responsible manner and shall handover the premises to the Lessor in working condition (reasonable wear and tear and loss or damage by fire, flood, rains, accidents, irresistible force or act of God excepted).")
    doc.add_paragraph(f"11. That in consideration of use of the said premises the Lessee agrees that he shall pay to the Lessor during the period of this agreement, a monthly rent at the rate of {answers[20]} ({answers[21]}). The amount will be paid in advance on or before the date of {answers[19]} of each month.")
    doc.add_paragraph(f"12. That the Lessee shall be entitled to, at his own cost, install and use electric gadgets and equipment such as air conditioners, electric heaters, geysers, coolers, refrigerators and other household equipments and facilities for which required power load is available. The Lessee shall be responsible for payment of the consumption charges of electricity in respect of the electric power consumed in the said premises. In case the Lessee does not replace the units with the requirements of the notice before its date of expiry, the Lessee shall be entitled to hold and retain possession of the said premises for the then unexpired portion of the said term.")
    doc.add_paragraph(f"13. That the Lessee shall deposit with the Lessor a sum of {answers[24]} ({answers[25]}) as a security deposit for due performance of this agreement and observance of the terms and conditions herein contained.")
    
    return doc

# Function to create trust deed document
def create_trust_deed(answers):
    doc = Document()
    
    doc.add_heading('Trust Deed', 0)

    doc.add_paragraph(f"This deed of trust is made on {answers[1]} by {answers[2]}, residing at {answers[3]}, {answers[4]}, {answers[5]}, hereinafter referred to as the 'Trustee'.")

    doc.add_paragraph(f"WHEREAS the Trustee is desirous of creating a trust in the name of {answers[0]} for the benefit of {answers[6]}, residing at {answers[7]}, {answers[8]}, {answers[9]}, hereinafter referred to as the 'Beneficiary';")

    doc.add_paragraph(f"AND WHEREAS the Trustee has agreed to contribute an initial capital amount of {answers[12]} ({answers[13]}) to the trust for achieving its objectives.")

    doc.add_paragraph(f"NOW THIS DEED WITNESSETH AS FOLLOWS:")

    doc.add_paragraph(f"1. The name of the trust shall be '{answers[0]}'.")
    doc.add_paragraph(f"2. The trust is created for the purpose of {answers[10]}.")
    doc.add_paragraph(f"3. The duration of the trust shall be {answers[11]}.")
    doc.add_paragraph(f"4. The initial capital amount of the trust shall be {answers[12]} ({answers[13]}).")
    doc.add_paragraph(f"5. The Trustee shall manage the trust funds and assets in accordance with the terms of this deed and for the benefit of the Beneficiary.")
    doc.add_paragraph(f"6. The Trustee shall keep proper accounts of the trust and provide regular updates to the Beneficiary.")

    return doc

# Function to create MoU document
def create_mou(answers):
    doc = Document()
    
    doc.add_heading('Memorandum of Understanding (MoU)', 0)

    doc.add_paragraph(f"This Memorandum of Understanding (MoU) is made on {answers[0]} between {answers[1]}, residing at {answers[2]}, {answers[3]}, {answers[4]}, hereinafter referred to as the 'First Party' and {answers[5]}, residing at {answers[6]}, {answers[7]}, {answers[8]}, hereinafter referred to as the 'Second Party'.")

    doc.add_paragraph(f"WHEREAS the First Party and the Second Party wish to collaborate on {answers[9]} for a duration of {answers[10]}.")

    doc.add_paragraph(f"NOW THEREFORE, the Parties agree to the following terms and conditions:")

    doc.add_paragraph(f"1. Purpose: The purpose of this MoU is to {answers[9]}.")
    doc.add_paragraph(f"2. Duration: This MoU shall remain in effect for {answers[10]}.")
    doc.add_paragraph(f"3. Terms and Conditions: {answers[11]}")

    return doc

# Streamlit UI for main page
st.title("Document Generator")

# Define the pages
PAGES = {
    "Rental Agreement": {
        "questions": [
            "Enter the city and state where the agreement is made:",
            "Enter the date of the agreement:",
            "Enter the landlord's name:",
            "Enter the landlord's address (Line 1):",
            "Enter the landlord's address (Line 2):",
            "Enter the landlord's city, state, and pin code:",
            "Enter the tenant's name:",
            "Enter the tenant's address (Line 1):",
            "Enter the tenant's address (Line 2):",
            "Enter the tenant's city, state, and pin code:",
            "Enter the description of the property:",
            "Enter the property address (Line 1):",
            "Enter the property address (Line 2):",
            "Enter the property city, state, and pin code:",
            "Enter the number of bedrooms, bathrooms, and carparks:",
            "Enter the monthly rent amount:",
            "Enter the security deposit amount:",
            "Enter the lease term in months:",
            "Enter the start date of the lease:",
            "Enter the notice period for termination:",
            "Enter the rent amount in words:",
            "Enter any additional clauses or requirements:",
            "Enter any other terms and conditions:",
        ],
        "function": create_rental_agreement
    },
    "Trust Deed": {
        "questions": [
            "Enter the name of the trust:",
            "Enter the date of the trust deed:",
            "Enter the trustee's name:",
            "Enter the trustee's address (Line 1):",
            "Enter the trustee's address (Line 2):",
            "Enter the trustee's city, state, and pin code:",
            "Enter the beneficiary's name:",
            "Enter the beneficiary's address (Line 1):",
            "Enter the beneficiary's address (Line 2):",
            "Enter the beneficiary's city, state, and pin code:",
            "Enter the purpose of the trust:",
            "Enter the duration of the trust:",
            "Enter the initial capital amount:",
            "Enter the capital amount in words:",
        ],
        "function": create_trust_deed
    },
    "MoU": {
        "questions": [
            "Enter the date of the MoU:",
            "Enter the first party's name:",
            "Enter the first party's address (Line 1):",
            "Enter the first party's address (Line 2):",
            "Enter the first party's city, state, and pin code:",
            "Enter the second party's name:",
            "Enter the second party's address (Line 1):",
            "Enter the second party's address (Line 2):",
            "Enter the second party's city, state, and pin code:",
            "Enter the purpose of the MoU:",
            "Enter the duration of the MoU:",
            "Enter the terms and conditions:",
        ],
        "function": create_mou
    },
}

# Sidebar for navigation
st.sidebar.title("Navigation")
selection = st.sidebar.radio("Go to", list(PAGES.keys()))

# Show the selected page
page = PAGES[selection]
st.header(f"{selection}")

# Get user input
answers = [st.text_input(q) for q in page["questions"]]

# Button to generate document
if st.button(f"Generate {selection}"):
    if all(answers):
        doc = page["function"](answers)
        file_path = f"{selection.replace(' ', '_')}.docx"
        doc.save(file_path)
        st.success(f"{selection} created successfully!")
        st.download_button(
            label=f"Download {selection}",
            data=open(file_path, "rb").read(),
            file_name=file_path,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error("Please fill in all the fields.")
