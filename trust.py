import streamlit as st
from docx import Document
from datetime import datetime

# Define the questions for the trust deed
trust_deed_questions = [
    "Enter the name of the trust:",
    "Enter the date of the trust deed:",
    "Enter the trustee's name:",
    "Enter the trustee's address (Line 1):",
    "Enter the trustee's address (Line 2):",
    "Enter the trustee's city, state, and pin code:",
    "Enter the beneficiary's name:",
    "Enter the beneficiary's address (Line 1):",
    "Enter the beneficiary's address (Line 2):",
    "Enter the beneficiary's city, state, and pin code:",
    "Enter the purpose of the trust:",
    "Enter the duration of the trust:",
    "Enter the initial capital amount of the trust (in numbers):",
    "Enter the initial capital amount of the trust (in words):"
]

def create_trust_deed(answers):
    doc = Document()
    
    doc.add_heading('Trust Deed', 0)

    doc.add_paragraph(f"This deed of trust is made on {answers[1]} by {answers[2]}, residing at {answers[3]}, {answers[4]}, {answers[5]}, hereinafter referred to as the 'Trustee'.")

    doc.add_paragraph(f"WHEREAS the Trustee is desirous of creating a trust in the name of {answers[0]} for the benefit of {answers[6]}, residing at {answers[7]}, {answers[8]}, {answers[9]}, hereinafter referred to as the 'Beneficiary';")

    doc.add_paragraph(f"AND WHEREAS the Trustee has agreed to contribute an initial capital amount of {answers[12]} ({answers[13]}) to the trust for achieving its objectives.")

    doc.add_paragraph(f"NOW THIS DEED WITNESSETH AS FOLLOWS:")

    doc.add_paragraph(f"1. The name of the trust shall be '{answers[0]}'.")
    doc.add_paragraph(f"2. The trust is created for the purpose of {answers[10]}.")
    doc.add_paragraph(f"3. The duration of the trust shall be {answers[11]}.")
    doc.add_paragraph(f"4. The initial capital amount of the trust shall be {answers[12]} ({answers[13]}).")
    doc.add_paragraph(f"5. The Trustee shall manage the trust funds and assets in accordance with the terms of this deed and for the benefit of the Beneficiary.")
    doc.add_paragraph(f"6. The Trustee shall keep proper accounts of the trust and provide regular updates to the Beneficiary.")

    return doc

# Streamlit UI for Trust Deed
st.title("Trust Deed Document Generator")
st.write("Answer the following questions to generate your trust deed:")

# Initialize session state to keep track of questions and answers
if 'current_trust_deed_question_index' not in st.session_state:
    st.session_state.current_trust_deed_question_index = 0
if 'trust_deed_answers' not in st.session_state:
    st.session_state.trust_deed_answers = {}

# Display the current question
if st.session_state.current_trust_deed_question_index < len(trust_deed_questions):
    question = trust_deed_questions[st.session_state.current_trust_deed_question_index]
    
    if "date" in question.lower():
        answer = st.date_input(question)
    else:
        answer = st.text_input(question)
    
    if st.button("Submit"):
        if answer:
            st.session_state.trust_deed_answers[st.session_state.current_trust_deed_question_index] = answer
            st.session_state.current_trust_deed_question_index += 1
            st.experimental_rerun()

    if st.button("Back") and st.session_state.current_trust_deed_question_index > 0:
        st.session_state.current_trust_deed_question_index -= 1
        st.experimental_rerun()
    
else:
    st.write("Thank you! Generating your trust deed...")

    answers = [st.session_state.trust_deed_answers.get(i, "") for i in range(len(trust_deed_questions))]
    document = create_trust_deed(answers)

    # Save the document
    document_path = "Trust_Deed.docx"
    document.save(document_path)

    st.write("Your trust deed has been generated.")
    with open(document_path, "rb") as file:
        st.download_button("Download Trust Deed", file, file_name=document_path)

    if st.button("Summary"):
        st.write("Summary of your inputs:")
        for i, answer in st.session_state.trust_deed_answers.items():
            st.write(f"{trust_deed_questions[i]} {answer}")
